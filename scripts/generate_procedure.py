# -*- coding: utf-8 -*-
#!/usr/bin/env python3
"""
CXskill - 程序文件生成脚本（docx 处理库，AI 驱动）
=====================================================
与 SCskill/generate_manual.py 结构一致，用于生成程序文件。
不同点：
  - find_template 查找"程序文件"分类而非"手册"
  - 全质知识库路径：external_kb/体系文件/程序文件/{部门}/
  - 企业内部文件路径：agent_{id}/程序文件/{部门}/
"""
import os
import sys
import json
import re
import subprocess
import tempfile
import shutil
import signal
import zipfile
from pathlib import Path
from datetime import date, datetime


def ensure_packages():
    import importlib
    try:
        importlib.import_module('docx')
    except ImportError:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install',
                               'python-docx', '--quiet'])


ensure_packages()

from docx import Document

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
TEMPLATES_DIR = SKILL_ROOT / "templates"


# ===================================================================
# 1. 模板查找（三级：企业内部→全质知识库→无内置兜底）
# ===================================================================

def find_template(agent_id=None, documents_dir=None):
    """查找程序文件模板（三级查找）：
    1. 企业内部文件知识库（agent_{agent_id}/程序文件/）递归搜索所有子目录
    2. 全质知识库（external_kb/体系文件/程序文件/）递归搜索所有子目录
    3. 无内置模板（程序文件无通用模板，必须从知识库找）

    返回 (template_path, need_convert, template_source, template_subdir)
    template_source: 'internal' / 'external' / None
    template_subdir: 模板所在的部门子目录名（如"总经办二级"）
    """
    if documents_dir is None:
        candidate = SKILL_ROOT.parent.parent / "data" / "documents"
        if candidate.exists():
            documents_dir = candidate
        else:
            documents_dir = SKILL_ROOT.parent / "data" / "documents"
    else:
        documents_dir = Path(documents_dir)

    print(f"[INFO] find_template: documents_dir={documents_dir}, agent_id={agent_id}")

    # 1. 企业内部文件知识库
    if agent_id:
        proc_dir = documents_dir / f"agent_{agent_id}" / "程序文件"
        print(f"[INFO] 查找企业内部文件: {proc_dir} (exists={proc_dir.exists()})")
        if proc_dir.exists():
            for root, dirs, files in os.walk(str(proc_dir)):
                for f in sorted(files):
                    if f.lower().endswith('.docx') and not f.startswith('~$'):
                        rel = os.path.relpath(root, str(proc_dir))
                        print(f"[INFO] 从企业内部文件知识库找到模板: {f} (路径: {root})")
                        return Path(root) / f, False, 'internal', rel
            for root, dirs, files in os.walk(str(proc_dir)):
                for f in sorted(files):
                    if f.lower().endswith('.doc') and not f.startswith('~$'):
                        rel = os.path.relpath(root, str(proc_dir))
                        print(f"[INFO] 从企业内部文件知识库找到 .doc 模板: {f} (路径: {root})")
                        return Path(root) / f, True, 'internal', rel
            print(f"[INFO] 企业内部文件知识库 程序文件/ 下未找到文件")

    # 2. 全质知识库
    ext_proc_dir = documents_dir / "external_kb" / "体系文件" / "程序文件"
    print(f"[INFO] 查找全质知识库: {ext_proc_dir} (exists={ext_proc_dir.exists()})")
    if ext_proc_dir.exists():
        for root, dirs, files in os.walk(str(ext_proc_dir)):
            for f in sorted(files):
                if f.lower().endswith('.docx') and not f.startswith('~$'):
                    rel = os.path.relpath(root, str(ext_proc_dir))
                    print(f"[INFO] 从全质知识库找到模板: {f} (路径: {root})")
                    return Path(root) / f, False, 'external', rel
        for root, dirs, files in os.walk(str(ext_proc_dir)):
            for f in sorted(files):
                if f.lower().endswith('.doc') and not f.startswith('~$'):
                    rel = os.path.relpath(root, str(ext_proc_dir))
                    print(f"[INFO] 从全质知识库找到 .doc 模板: {f} (路径: {root})")
                    return Path(root) / f, True, 'external', rel

    # 3. 无内置模板
    print(f"[INFO] 未找到任何程序文件模板")
    return None, False, None, None


def find_all_templates(agent_id=None, documents_dir=None):
    """查找所有程序文件模板（收集所有部门的所有文件）

    查找逻辑：
    1. 企业内部文件 agent_{id}/程序文件/ → 递归搜索所有子目录和文件
    2. 全质知识库 external_kb/体系文件/程序文件/ → 递归搜索所有子目录和文件
    3. 按部门去重：如果企业内部文件有某部门的模板，全质知识库的同部门模板跳过

    返回 list of dict:
    [
      {"path": Path, "need_convert": bool, "source": "internal"/"external",
       "dept": "部门名", "filename": "文件名"},
      ...
    ]
    """
    if documents_dir is None:
        candidate = SKILL_ROOT.parent.parent / "data" / "documents"
        if candidate.exists():
            documents_dir = candidate
        else:
            documents_dir = SKILL_ROOT.parent / "data" / "documents"
    else:
        documents_dir = Path(documents_dir)

    print(f"[INFO] find_all_templates: documents_dir={documents_dir}, agent_id={agent_id}")

    # 收集企业内部文件
    internal_templates = []
    if agent_id:
        proc_dir = documents_dir / f"agent_{agent_id}" / "程序文件"
        print(f"[INFO] 查找企业内部文件: {proc_dir} (exists={proc_dir.exists()})")
        if proc_dir.exists():
            for root, dirs, files in os.walk(str(proc_dir)):
                for f in sorted(files):
                    if f.startswith('~$'):
                        continue
                    ext = f.lower().rsplit('.', 1)[-1] if '.' in f else ''
                    if ext in ('docx', 'doc', 'xlsx'):
                        rel = os.path.relpath(root, str(proc_dir))
                        # 部门名取第一级子目录（如"质量部二级"），如果文件直接在根目录则 dept="通用"
                        parts = rel.replace('\\', '/').split('/')
                        dept = parts[0] if parts and parts[0] != '.' else '通用'
                        internal_templates.append({
                            "path": Path(root) / f,
                            "need_convert": (ext == 'doc'),
                            "file_type": ext,
                            "source": 'internal',
                            "dept": dept,
                            "filename": f
                        })
                        print(f"[INFO] 企业内部文件: {dept}/{f} (source=internal)")

    # 收集全质知识库
    ext_templates = []
    ext_proc_dir = documents_dir / "external_kb" / "体系文件" / "程序文件"
    print(f"[INFO] 查找全质知识库: {ext_proc_dir} (exists={ext_proc_dir.exists()})")
    if ext_proc_dir.exists():
        for root, dirs, files in os.walk(str(ext_proc_dir)):
            for f in sorted(files):
                if f.startswith('~$'):
                    continue
                ext = f.lower().rsplit('.', 1)[-1] if '.' in f else ''
                if ext in ('docx', 'doc'):
                    # [需求] 全质知识库模板只使用 AAA/aaa 命名的文件
                    # AAB、AAC 等是其他公司的模板，不用于生成
                    if 'AAA' not in f and 'aaa' not in f:
                        print(f"[INFO] 跳过非AAA模板: {f}")
                        continue
                    rel = os.path.relpath(root, str(ext_proc_dir))
                    parts = rel.replace('\\', '/').split('/')
                    dept = parts[0] if parts and parts[0] != '.' else '通用'
                    ext_templates.append({
                        "path": Path(root) / f,
                        "need_convert": (ext == 'doc'),
                        "file_type": ext,
                        "source": 'external',
                        "dept": dept,
                        "filename": f
                    })
                    print(f"[INFO] 全质知识库: {dept}/{f} (source=external)")

    # 不再按部门整体去重，改为返回所有模板
    # 同部门去重逻辑交给 API 层用 AI 智能判断（按文件内容主题匹配，而非按部门整体跳过）
    all_templates = internal_templates + ext_templates
    print(f"[INFO] 总计找到 {len(all_templates)} 个模板（企业内部 {len(internal_templates)} + 全质知识库 {len(ext_templates)}）")
    for t in all_templates:
        print(f"  - [{t['source']}] {t['dept']}/{t['filename']}")

    return all_templates


def _find_libreoffice_executable():
    """返回一个可执行的 LibreOffice soffice 路径，未安装时返回 None。

    不直接调用字符串形式的 ``soffice``，以免在 Windows 上对每个文件产生
    ``WinError 2`` 日志；也不猜测或终止其他请求正在使用的 LibreOffice 进程。
    """
    candidates = []
    configured = os.environ.get('LIBREOFFICE_PATH')
    if configured:
        candidates.append(configured)

    for command in ('soffice', 'libreoffice'):
        resolved = shutil.which(command)
        if resolved:
            candidates.append(resolved)

    if os.name == 'nt':
        program_files = [
            os.environ.get('ProgramFiles'),
            os.environ.get('ProgramFiles(x86)'),
            r'C:\Program Files',
            r'C:\Program Files (x86)',
        ]
        for root in program_files:
            if root:
                candidates.append(os.path.join(root, 'LibreOffice', 'program', 'soffice.exe'))
    else:
        candidates.extend(('/usr/bin/soffice', '/usr/local/bin/soffice'))

    seen = set()
    for candidate in candidates:
        candidate = os.path.expandvars(os.path.expanduser(candidate))
        normalized = os.path.normcase(os.path.abspath(candidate))
        if normalized in seen:
            continue
        seen.add(normalized)
        if os.path.isfile(candidate):
            return candidate
    return None


def _cleanup_conversion_temp_dir(tmp_dir):
    """尽力清理一次失败转换留下的输出目录和独立 LO 配置目录。"""
    try:
        shutil.rmtree(tmp_dir)
    except FileNotFoundError:
        pass
    except Exception as exc:
        # 转换结果不可用时不应让清理异常掩盖原始转换错误。
        print(f"[WARN] 转换临时目录清理失败 ({tmp_dir}): {exc}")


def _convert_doc_with_word(doc_path, tmp_dir):
    """在 Windows 上使用独立 Word COM 实例进行兜底转换。"""
    word = None
    source_doc = None
    pythoncom = None
    com_initialized = False
    output_path = os.path.join(tmp_dir, 'converted.docx')

    try:
        import win32com.client
        import pythoncom as _pythoncom

        pythoncom = _pythoncom
        pythoncom.CoInitialize()
        com_initialized = True

        # DispatchEx 创建本次转换专用实例，避免关闭已有用户打开的 Word。
        dispatch = getattr(win32com.client, 'DispatchEx', win32com.client.Dispatch)
        word = dispatch('Word.Application')
        word.Visible = False
        try:
            word.DisplayAlerts = 0
        except Exception:
            pass

        source_doc = word.Documents.Open(str(Path(doc_path).resolve()), ReadOnly=True)
        source_doc.SaveAs2(output_path, FileFormat=16)
        if os.path.exists(output_path):
            print("[INFO] Word COM 转换成功")
            return output_path
        print("[WARN] Word COM 未生成转换后的 .docx 文件")
    except Exception as exc:
        print(f"[WARN] Word COM 转换失败: {exc}")
    finally:
        if source_doc is not None:
            try:
                source_doc.Close(False)
            except Exception:
                try:
                    source_doc.Close()
                except Exception:
                    pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if com_initialized and pythoncom is not None:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    return None


def _legacy_convert_doc_to_docx(doc_path):
    """用 LibreOffice 把 .doc 转成 .docx，失败时在 Windows 上回退 Word COM。

    每次转换均使用一个临时、独立的 LibreOffice 用户配置目录，避免共享 profile
    锁造成卡死。不会再 ``taskkill``/``pkill`` 全局 soffice 进程，也不会进行会把
    其他并发转换杀掉的重试。成功时保留返回文件所在临时目录供调用方继续读取；
    失败时会在返回 ``None`` 前清理该目录。
    """
    tmp_dir = tempfile.mkdtemp(prefix='doc2docx_')
    converted = None

    try:
        soffice = _find_libreoffice_executable()
        if soffice:
            profile_dir = Path(tmp_dir) / 'libreoffice-profile'
            profile_dir.mkdir(parents=True, exist_ok=True)
            profile_uri = profile_dir.resolve().as_uri()
            expected_output = os.path.join(
                tmp_dir, Path(doc_path).stem + '.docx'
            )
            try:
                result = subprocess.run(
                    [
                        soffice,
                        f'-env:UserInstallation={profile_uri}',
                        '--headless', '--convert-to', 'docx',
                        '--outdir', tmp_dir, str(doc_path),
                    ],
                    capture_output=True,
                    text=True,
                    errors='replace',
                    timeout=120,
                )
                if result.returncode == 0 and os.path.exists(expected_output):
                    print('[INFO] LibreOffice 转换成功')
                    converted = expected_output
                elif result.returncode == 0:
                    print('[WARN] LibreOffice 未生成预期的 .docx 文件')
                else:
                    stderr = (result.stderr or '').strip().replace('\n', ' ')
                    print(f'[WARN] LibreOffice ({soffice}) 返回码 {result.returncode}: {stderr[:200]}')
            except subprocess.TimeoutExpired:
                # subprocess.run 只终止当前启动的进程，不会影响其他用户的转换。
                print(f'[WARN] LibreOffice ({soffice}) 转换超时（120s），将尝试 Word COM 兜底')
            except Exception as exc:
                print(f'[WARN] LibreOffice ({soffice}) 转换失败: {exc}')
        else:
            print('[WARN] 未找到 LibreOffice soffice 可执行文件，将尝试 Word COM 兜底')

        if converted is None:
            converted = _convert_doc_with_word(doc_path, tmp_dir)
        return converted
    finally:
        if converted is None:
            _cleanup_conversion_temp_dir(tmp_dir)


# ===================================================================
# 2. 提取模板结构概览
# ===================================================================

def _find_libreoffice_executables():
    """返回当前机器真实存在的 LibreOffice 可执行文件，按路径去重。"""
    candidates = []
    configured = os.environ.get('LIBREOFFICE_PATH')
    if configured:
        candidates.append(configured)
    for command in ('soffice', 'libreoffice'):
        resolved = shutil.which(command)
        if resolved:
            candidates.append(resolved)
    if os.name == 'nt':
        for root in (
            os.environ.get('ProgramW6432'),
            os.environ.get('ProgramFiles'),
            os.environ.get('ProgramFiles(x86)'),
        ):
            if root:
                candidates.append(str(Path(root) / 'LibreOffice' / 'program' / 'soffice.exe'))
    else:
        candidates.extend((
            '/usr/bin/soffice', '/usr/local/bin/soffice',
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        ))
    executables = []
    seen = set()
    for candidate in candidates:
        try:
            normalized = str(Path(candidate).expanduser().resolve())
        except (OSError, TypeError, ValueError):
            continue
        key = os.path.normcase(normalized)
        if key not in seen and Path(normalized).is_file():
            seen.add(key)
            executables.append(normalized)
    return executables


def _terminate_process_tree(process):
    """只终止本次转换启动的进程树，不影响其他用户的 LibreOffice。"""
    if process is None or process.poll() is not None:
        return
    try:
        if os.name == 'nt':
            subprocess.run(
                ['taskkill', '/PID', str(process.pid), '/T', '/F'],
                capture_output=True, timeout=10,
            )
        else:
            os.killpg(process.pid, signal.SIGKILL)
    except Exception:
        try:
            process.kill()
        except Exception:
            pass


def _run_conversion_process(command, timeout):
    """运行一次可控转换；超时时回收本次 PID 及其子进程。"""
    kwargs = {
        'stdout': subprocess.PIPE, 'stderr': subprocess.PIPE,
        'text': True, 'errors': 'replace',
    }
    if os.name == 'nt':
        kwargs['creationflags'] = getattr(subprocess, 'CREATE_NEW_PROCESS_GROUP', 0)
    else:
        kwargs['start_new_session'] = True
    process = subprocess.Popen(command, **kwargs)
    try:
        stdout, stderr = process.communicate(timeout=timeout)
        return subprocess.CompletedProcess(command, process.returncode, stdout, stderr)
    except subprocess.TimeoutExpired:
        _terminate_process_tree(process)
        try:
            process.communicate(timeout=5)
        except Exception:
            pass
        raise


def _validate_docx(path):
    """确认产物是包含 Word 主文档的有效 OOXML 文件。"""
    candidate = Path(path)
    if not candidate.is_file() or candidate.stat().st_size <= 0:
        return False
    try:
        with zipfile.ZipFile(candidate, 'r') as archive:
            return 'word/document.xml' in archive.namelist()
    except (OSError, zipfile.BadZipFile):
        return False


def _convert_with_libreoffice(
    soffice_path, source_path, output_dir, output_ext='docx',
    input_filter=None, timeout=75, attempt_name='direct',
):
    """用独立 profile、英文临时路径和明确过滤器执行一次转换。"""
    source_path = Path(source_path)
    output_dir = Path(output_dir)
    profile_dir = output_dir / f'libreoffice-profile-{attempt_name}'
    profile_dir.mkdir(parents=True, exist_ok=True)
    converted = output_dir / f'{source_path.stem}.{output_ext}'
    try:
        converted.unlink(missing_ok=True)
    except OSError:
        pass
    export_filters = {'docx': 'Office Open XML Text', 'rtf': 'Rich Text Format'}
    convert_spec = output_ext
    if output_ext in export_filters:
        convert_spec += ':' + export_filters[output_ext]
    command = [
        str(soffice_path),
        f'-env:UserInstallation={profile_dir.resolve().as_uri()}',
        '--headless', '--invisible', '--nologo', '--nodefault',
        '--nofirststartwizard', '--norestore',
    ]
    if input_filter:
        command.append(f'--infilter={input_filter}')
    command.extend([
        '--convert-to', convert_spec, '--outdir', str(output_dir), str(source_path),
    ])
    try:
        result = _run_conversion_process(command, timeout)
        if result.returncode == 0 and converted.is_file() and converted.stat().st_size > 0:
            print(f'[INFO] LibreOffice 转换成功（{attempt_name}）')
            return str(converted)
        detail = (result.stderr or result.stdout or '').strip().replace('\n', ' ')
        print(f'[WARN] LibreOffice ({soffice_path}) {attempt_name} 失败，返回码 {result.returncode}: {detail[:300]}')
    except subprocess.TimeoutExpired:
        print(f'[WARN] LibreOffice ({soffice_path}) {attempt_name} 超时（{timeout}s）')
    except Exception as exc:
        print(f'[WARN] LibreOffice ({soffice_path}) {attempt_name} 失败: {exc}')
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)
    return None


def _convert_via_rtf(soffice_path, source_path, output_dir):
    """直接 DOCX 失败时，以 RTF 为中间格式修复旧 DOC 结构。"""
    rtf_path = _convert_with_libreoffice(
        soffice_path, source_path, output_dir,
        output_ext='rtf', input_filter='MS Word 97', timeout=45,
        attempt_name='doc-to-rtf',
    )
    if not rtf_path:
        return None
    return _convert_with_libreoffice(
        soffice_path, rtf_path, output_dir,
        output_ext='docx', input_filter='Rich Text Format', timeout=45,
        attempt_name='rtf-to-docx',
    )


def _convert_with_word_com(doc_path, tmp_dir):
    """Windows Word COM 最后兜底；使用独立实例并验证最终 DOCX。"""
    pythoncom = None
    word = None
    source_doc = None
    initialized = False
    try:
        import pythoncom as _pythoncom
        import win32com.client
        pythoncom = _pythoncom
        pythoncom.CoInitialize()
        initialized = True
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0
        source_doc = word.Documents.Open(str(doc_path), ReadOnly=True, AddToRecentFiles=False)
        output = Path(tmp_dir) / 'converted.docx'
        source_doc.SaveAs2(str(output), FileFormat=16)
        if output.is_file() and output.stat().st_size > 0:
            print('[INFO] Word COM 转换成功')
            return str(output)
    except Exception as exc:
        print(f'[WARN] Word COM 转换失败: {exc}')
    finally:
        if source_doc is not None:
            try:
                source_doc.Close(SaveChanges=0)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if initialized and pythoncom is not None:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
    return None


def convert_doc_to_docx(doc_path):
    """将旧 DOC 可靠转换为 DOCX；成功返回临时路径，失败返回 None。"""
    source = Path(doc_path)
    tmp_dir = Path(tempfile.mkdtemp(prefix='doc2docx_'))
    succeeded = False
    try:
        if not source.is_file():
            print(f'[WARN] .doc 文件不存在，无法转换: {source}')
            return None
        # 旧格式过滤器在 Windows 上可能被中文/特殊字符路径卡住。
        safe_source = tmp_dir / 'legacy_input.doc'
        shutil.copy2(source, safe_source)
        soffice_paths = _find_libreoffice_executables()
        for soffice_path in soffice_paths:
            converted = _convert_with_libreoffice(
                soffice_path, safe_source, tmp_dir,
                output_ext='docx', input_filter='MS Word 97', timeout=75,
                attempt_name='word97-direct',
            )
            if converted and _validate_docx(converted):
                succeeded = True
                return converted
            converted = _convert_via_rtf(soffice_path, safe_source, tmp_dir)
            if converted and _validate_docx(converted):
                print('[INFO] 旧 DOC 已通过 RTF 修复链转换为 DOCX')
                succeeded = True
                return converted
        if not soffice_paths:
            print('[WARN] 未找到可用的 LibreOffice，将尝试 Word COM')
        converted = _convert_with_word_com(safe_source, tmp_dir)
        if converted and _validate_docx(converted):
            succeeded = True
            return converted
        if converted:
            print('[WARN] Word COM 生成的文件不是有效 DOCX，已丢弃')
        return None
    finally:
        if not succeeded:
            shutil.rmtree(tmp_dir, ignore_errors=True)


def load_xlsx_workbook(xlsx_path):
    """加载 XLSX 模板；保留公式，以便只修改输入数据。"""
    from openpyxl import load_workbook
    return load_workbook(str(xlsx_path), data_only=False)


def _is_formula_cell(cell):
    return cell.data_type == 'f' or (
        isinstance(cell.value, str) and cell.value.startswith('=')
    )


def _iter_existing_xlsx_cells(worksheet):
    """只遍历工作簿中真实存在的单元格，跳过被整列格式撑大的空白区域。"""
    existing = getattr(worksheet, '_cells', None)
    if isinstance(existing, dict):
        return sorted(existing.values(), key=lambda cell: (cell.row, cell.column))
    return (cell for row in worksheet.iter_rows() for cell in row)


def build_xlsx_overview_chunks(
    workbook, max_cells=160, max_cell_chars=12000, max_chunk_chars=40000,
):
    """分块返回所有非空、非公式单元格；不截断工作簿后半部分。"""
    entries = []
    for worksheet in workbook.worksheets:
        for cell in _iter_existing_xlsx_cells(worksheet):
            value = cell.value
            if value is None or value == '' or _is_formula_cell(cell):
                continue
            if isinstance(value, datetime):
                display_value = value.isoformat(sep=' ')
                value_type = 'date'
            else:
                display_value = str(value)
                value_type = cell.data_type or type(value).__name__
            if len(display_value) > max_cell_chars:
                half = max_cell_chars // 2
                display_value = display_value[:half] + '\n...[内容过长，中间省略]...\n' + display_value[-half:]
            entries.append({
                'sheet': worksheet.title,
                'cell': cell.coordinate,
                'value_type': value_type,
                'value': display_value,
            })
    chunks = []
    lines = []
    chars = 0
    for item in entries:
        line = json.dumps(item, ensure_ascii=False, separators=(',', ':'))
        if lines and (len(lines) >= max_cells or chars + len(line) + 1 > max_chunk_chars):
            chunks.append('\n'.join(lines))
            lines = []
            chars = 0
        lines.append(line)
        chars += len(line) + 1
    if lines:
        chunks.append('\n'.join(lines))
    return chunks


def build_xlsx_llm_prompt(
    overview_text, survey_text, template_filename, survey=None,
):
    """构造程序文件 XLSX 的逐单元格 AI 检查提示词。"""
    system = (
        "你是一名 IATF 16949/ISO 9001 体系程序文件专家。请结合企业体系调研信息，"
        "逐个检查 Excel 模板中本批次列出的每一个单元格，只对确实需要企业化的内容提出修改。\n\n"
        "必须检查：公司名称、文件编号/代号、人员姓名、职务、部门名称、日期、地址、电话、"
        "客户名称以及明显属于示例企业的占位内容。表格内容少也必须检查。\n"
        "禁止修改：标准条款本身、通用制度要求、公式、工作表名称、行列结构、合并单元格和样式。"
        "调研信息没有提供的字段不得臆造；new_value 必须保持原单元格的数据类型，日期使用 YYYY-MM-DD。\n\n"
        "每行只输出一个 JSON 对象，可用格式：\n"
        '{"type":"xlsx_cell","sheet":"工作表名","cell":"B2",'
        '"new_value":"新值","reason":"原因"}\n'
        '{"type":"xlsx_global_replace","old":"明确的旧占位文本",'
        '"new":"新值","reason":"原因"}\n'
        "xlsx_global_replace 只适合确定应在整个工作簿统一替换的完整旧文本；其余情况使用 xlsx_cell。\n"
        "不得输出 markdown、解释或数组。没有修改时只输出 ===END===；最后一行必须输出 ===END===。"
    )
    user = (
        f"文件名：{template_filename}\n\n"
        f"=== 企业体系调研信息 ===\n{survey_text}\n\n"
        f"=== 本批次 Excel 单元格（JSON 行） ===\n{overview_text}\n\n"
        "请逐项核对后输出修改方案。"
    )
    return system, user


def _xlsx_merged_anchor(worksheet, coordinate):
    for merged_range in worksheet.merged_cells.ranges:
        if coordinate in merged_range:
            return worksheet.cell(
                merged_range.min_row, merged_range.min_col
            ).coordinate
    return coordinate


def _coerce_xlsx_value(old_value, new_value):
    """尽量维持原单元格数据类型，避免日期/数字被保存成普通文本。"""
    if not isinstance(new_value, str):
        return new_value
    text = new_value.strip()
    if isinstance(old_value, datetime):
        for candidate in (text, text.replace('/', '-')):
            try:
                return datetime.fromisoformat(candidate)
            except ValueError:
                continue
    elif isinstance(old_value, date):
        try:
            return date.fromisoformat(text.replace('/', '-'))
        except ValueError:
            pass
    elif isinstance(old_value, bool):
        lowered = text.lower()
        if lowered in ('true', 'yes', '1', '是'):
            return True
        if lowered in ('false', 'no', '0', '否'):
            return False
    elif isinstance(old_value, int):
        try:
            return int(text)
        except ValueError:
            pass
    elif isinstance(old_value, float):
        try:
            return float(text)
        except ValueError:
            pass
    return new_value


def apply_xlsx_cell_replace(workbook, sheet_name, coordinate, new_value):
    """更新一个非公式单元格的值，同时保留样式与数字格式。"""
    if sheet_name not in workbook.sheetnames:
        print(f"[WARN] XLSX 工作表不存在: {sheet_name}")
        return False
    worksheet = workbook[sheet_name]
    try:
        anchor = _xlsx_merged_anchor(worksheet, coordinate)
        if anchor != coordinate:
            print(f"[WARN] XLSX 单元格 {sheet_name}!{coordinate} 是合并区域非锚点，已跳过")
            return False
        cell = worksheet[coordinate]
    except (KeyError, TypeError, ValueError):
        print(f"[WARN] XLSX 单元格地址无效: {sheet_name}!{coordinate}")
        return False
    if _is_formula_cell(cell):
        print(f"[WARN] XLSX 公式单元格禁止修改: {sheet_name}!{coordinate}")
        return False
    new_value = _coerce_xlsx_value(cell.value, new_value)
    if cell.value == new_value:
        return False
    cell.value = new_value
    return True


def apply_xlsx_global_replace(workbook, old, new):
    """替换文本单元格里的明确旧文本；公式和非字符串值不受影响。"""
    if not isinstance(old, str) or not old or new is None or old == new:
        return 0
    count = 0
    for worksheet in workbook.worksheets:
        for cell in _iter_existing_xlsx_cells(worksheet):
            if _is_formula_cell(cell) or not isinstance(cell.value, str):
                continue
            if old in cell.value:
                cell.value = cell.value.replace(old, str(new))
                count += 1
    return count


def apply_xlsx_modifications(workbook, modifications):
    """全局替换优先，随后以精确单元格修改覆盖。"""
    stats = {
        'xlsx_cell': 0, 'xlsx_global_replace': 0,
        'unknown': 0, 'failed': 0,
    }
    globals_seen = set()
    global_modifications = []
    cell_seen = set()
    cell_modifications = []
    for modification in modifications or []:
        mod_type = modification.get('type')
        if mod_type == 'xlsx_global_replace':
            old = modification.get('old')
            new = modification.get('new')
            key = (str(old), json.dumps(new, ensure_ascii=False, sort_keys=True))
            if old and key not in globals_seen:
                globals_seen.add(key)
                global_modifications.append(modification)
        elif mod_type == 'xlsx_cell':
            sheet = modification.get('sheet')
            coordinate = str(modification.get('cell') or '').upper()
            key = (sheet, coordinate)
            if sheet and coordinate and key not in cell_seen:
                cell_seen.add(key)
                normalized = dict(modification)
                normalized['cell'] = coordinate
                cell_modifications.append(normalized)
        else:
            stats['unknown'] += 1
    global_modifications.sort(key=lambda item: len(str(item.get('old') or '')), reverse=True)
    for modification in global_modifications:
        try:
            stats['xlsx_global_replace'] += apply_xlsx_global_replace(
                workbook, modification.get('old'), modification.get('new'),
            )
        except Exception as exc:
            stats['failed'] += 1
            print(f"[WARN] XLSX 全局替换失败: {exc}")
    for modification in cell_modifications:
        try:
            if apply_xlsx_cell_replace(
                workbook, modification.get('sheet'), modification.get('cell'),
                modification.get('new_value'),
            ):
                stats['xlsx_cell'] += 1
        except Exception as exc:
            stats['failed'] += 1
            print(f"[WARN] XLSX 单元格修改失败: {exc}")
    return stats


def save_xlsx_workbook(workbook, output_path):
    """保存 XLSX，并要求 Excel 下次打开时重算公式。"""
    calculation = getattr(workbook, 'calculation', None)
    if calculation is not None:
        for name, value in (
            ('fullCalcOnLoad', True), ('forceFullCalc', True), ('calcMode', 'auto'),
        ):
            try:
                setattr(calculation, name, value)
            except Exception:
                pass
    workbook.save(str(output_path))
    return str(output_path)


def extract_template_overview(doc, max_paras=None):
    """提取模板的结构化概览"""
    overview = {"paragraphs": [], "tables": [], "headers": [], "footers": []}

    for i, p in enumerate(doc.paragraphs):
        if max_paras and i >= max_paras:
            break
        text = p.text.strip()
        if not text:
            continue
        try:
            style = p.style.name if p.style else ''
        except Exception:
            style = ''
        if 'toc' in (style or '').lower():
            continue
        overview["paragraphs"].append({
            "index": i, "style": style, "text": text[:500]
        })

    for ti, t in enumerate(doc.tables):
        rows = []
        for row in t.rows:
            cells = []
            for c in row.cells:
                cell_text = c.text.strip().replace('\n', ' | ')
                cells.append(cell_text[:200])
            rows.append(cells)
        overview["tables"].append({"index": ti, "rows": rows})

    for si, sec in enumerate(doc.sections):
        for hf_name in ['header', 'first_page_header', 'even_page_header']:
            hf = getattr(sec, hf_name, None)
            if not hf:
                continue
            texts = []
            for p in hf.paragraphs:
                if p.text.strip():
                    texts.append(p.text.strip())
            for t in hf.tables:
                for row in t.rows:
                    for c in row.cells:
                        if c.text.strip():
                            texts.append(c.text.strip().replace('\n', ' | '))
            if texts:
                overview["headers"].append({
                    "section": si, "type": hf_name, "text": " | ".join(texts)[:500]
                })
        for hf_name in ['footer', 'first_page_footer', 'even_page_footer']:
            hf = getattr(sec, hf_name, None)
            if not hf:
                continue
            texts = []
            for p in hf.paragraphs:
                if p.text.strip():
                    texts.append(p.text.strip())
            # 页脚中的文件编号、公司名和版次通常位于表格内。若不把它们送入
            # LLM 概览，模型既看不到内容，也无法生成 header_replace 方案。
            for t in hf.tables:
                for row in t.rows:
                    for c in row.cells:
                        if c.text.strip():
                            texts.append(c.text.strip().replace('\n', ' | '))
            if texts:
                overview["footers"].append({
                    "section": si, "type": hf_name, "text": " | ".join(texts)[:300]
                })

    return overview


def format_overview_for_llm(overview):
    lines = []
    lines.append("=== 模板段落（非空，已跳过目录） ===")
    for p in overview["paragraphs"]:
        lines.append(f"[P{p['index']}] ({p['style']}) {p['text']}")
    lines.append("")
    lines.append("=== 模板表格 ===")
    for t in overview["tables"]:
        lines.append(f"--- Table {t['index']} ({len(t['rows'])} rows) ---")
        for ri, row in enumerate(t['rows']):
            lines.append(f"  T{t['index']}.R{ri}: {row}")
    lines.append("")
    lines.append("=== 页眉 ===")
    for h in overview["headers"]:
        lines.append(f"[H{h['section']}.{h['type']}] {h['text']}")
    lines.append("")
    lines.append("=== 页脚 ===")
    for f in overview["footers"]:
        lines.append(f"[F{f['section']}.{f['type']}] {f['text']}")
    return "\n".join(lines)


def format_survey_for_llm(survey):
    # [修复] 字段列表与 SCskill/generate_manual.py 完全对齐，
    # 确保程序文件生成时 AI 能看到所有调研字段（含认证计划日期、贯标组、内审员等）
    field_labels = {
        'sv_company_name': '公司名称',
        'sv_cert_other': '其他证书',
        'sv_chairman': '董事长',
        'sv_legal_rep': '法人代表',
        'sv_gm': '总经理',
        'sv_deputy_gm': '副总经理',
        'sv_mgmt_rep': '管理者代表',
        'sv_leader_group_leader': '贯标组长',
        'sv_leader_group_members': '贯标组员',
        'sv_iso_office_head': '贯标办主任',
        'sv_iso_office_members': '贯标办成员',
        'sv_auditors': '内审员',
        'sv_products': '体系覆盖产品',
        'sv_process_flow': '生产流程',
        'sv_location': '地理位置',
        'sv_area': '占地面积',
        'sv_building_area': '建筑面积',
        'sv_staff_total': '正式员工人数',
        'sv_staff_mgmt': '管理技术人员数',
        'sv_staff_edu': '学历分布',
        'sv_equipment': '设备情况',
        'sv_customers': '主要客户',
        'sv_address': '公司地址',
        'sv_contact': '联系人',
        'sv_phone': '电话',
        'sv_fax': '传真',
        'sv_mobile': '手机',
        'sv_purpose': '公司宗旨/经营理念',
        'sv_quality_policy': '质量方针',
        'sv_quality_goal': '质量目标',
        'sv_cert_date': '认证日期',
        'sv_audit_date': '审核日期',
        'sv_rest_day': '休息日',
        'sv_design_dev': '有无设计开发',
        'sv_filler_name': '填写人',
        'sv_filler_phone': '填写人手机',
        'sv_certs': '已有证书',
        'sv_org': '机构设置',
        'sv_org_custom_rows': '自定义部门',
    }
    lines = ["=== 用户体系调研数据 ==="]
    for key, label in field_labels.items():
        val = survey.get(key, '')
        if isinstance(val, (list, dict)):
            val = json.dumps(val, ensure_ascii=False)
        if val:
            lines.append(f"{label}（{key}）：{val}")
        else:
            lines.append(f"{label}（{key}）：[未填写]")
    return "\n".join(lines)


# ===================================================================
# 2.5 实施日期计算（与 SCskill/generate_manual.py 逻辑完全一致）
# ===================================================================

def _parse_date(date_str):
    """将各种格式的日期字符串解析为 datetime 对象，失败返回 None。
    支持: 2024年3月15日, 2024-03-15, 2024/03/15, 2024.03.15, ISO格式"""
    if not date_str or not isinstance(date_str, str):
        return None
    date_str = date_str.strip()
    if not date_str:
        return None
    formats = [
        '%Y年%m月%d日', '%Y-%m-%d', '%Y/%m/%d', '%Y.%m.%d',
        '%Y-%m-%dT%H:%M:%S', '%Y-%m-%dT%H:%M:%S.%f',
        '%Y-%m-%d %H:%M:%S',
        '%Y年%m月%d日 %H:%M:%S',
    ]
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    m = re.search(r'(\d{4})\D+(\d{1,2})\D+(\d{1,2})', date_str)
    if m:
        try:
            return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
        except ValueError:
            return None
    return None


def _contains_16949(certs):
    """判断证书列表中是否包含 IATF 16949"""
    if not certs:
        return False
    if isinstance(certs, str):
        return '16949' in certs
    if isinstance(certs, (list, tuple)):
        for c in certs:
            if isinstance(c, str) and '16949' in c:
                return True
    return False


def _subtract_months(dt, months):
    """从日期减去指定月数，返回新的 datetime。自动处理月末溢出（如 3月31日 - 1月 = 2月28/29日）。"""
    import calendar
    total_months = dt.year * 12 + dt.month - 1 - months
    new_year = total_months // 12
    new_month = total_months % 12 + 1
    max_day = calendar.monthrange(new_year, new_month)[1]
    new_day = min(dt.day, max_day)
    return dt.replace(year=new_year, month=new_month, day=new_day)


def calculate_implementation_date(survey):
    """根据证书类型和第一阶段开始时间，计算程序文件实施日期（与手册保持一致）。

    逻辑（与 SCskill 完全一致，确保体系文件时间基准统一）：
    - 证书含 IATF 16949:
      - 未填写第一阶段开始时间 → 计划取得认证证书日期 - 18个月
      - 已填写第一阶段开始时间 → 第一阶段开始时间 - 15个月
    - 非 16949（如 ISO 9001）:
      - 未填写第一阶段开始时间 → 计划取得认证证书日期 - 5个月
      - 已填写第一阶段开始时间 → 第一阶段开始时间 - 4个月

    返回 (implementation_date_str, description_str)
    implementation_date_str 格式: "2024年3月15日"
    """
    certs = survey.get('sv_certs', [])
    is_16949 = _contains_16949(certs)

    audit_date_str = (survey.get('sv_audit_date') or '').strip()
    cert_date_str = (survey.get('sv_cert_date') or '').strip()

    audit_date = _parse_date(audit_date_str)
    cert_date = _parse_date(cert_date_str)

    if is_16949:
        if audit_date:
            impl_date = _subtract_months(audit_date, 15)
            desc = f"IATF 16949认证，第一阶段开始时间={audit_date_str}，向前推15个月作为程序文件实施日期"
        elif cert_date:
            impl_date = _subtract_months(cert_date, 18)
            desc = f"IATF 16949认证，计划取得认证证书日期={cert_date_str}，向前推18个月作为程序文件实施日期"
        else:
            today = datetime.now()
            impl_date = today
            desc = "IATF 16949认证，但认证日期和第一阶段开始时间均未填写，回退使用当前日期"
    else:
        if audit_date:
            impl_date = _subtract_months(audit_date, 4)
            desc = f"非16949认证，第一阶段开始时间={audit_date_str}，向前推4个月作为程序文件实施日期"
        elif cert_date:
            impl_date = _subtract_months(cert_date, 5)
            desc = f"非16949认证，计划取得认证证书日期={cert_date_str}，向前推5个月作为程序文件实施日期"
        else:
            today = datetime.now()
            impl_date = today
            desc = "非16949认证，但认证日期和第一阶段开始时间均未填写，回退使用当前日期"

    impl_date_str = f"{impl_date.year}年{impl_date.month}月{impl_date.day}日"
    return impl_date_str, desc


# ===================================================================
# 3. 构造 LLM 提示词（NDJSON 流式输出）
# ===================================================================

def build_llm_prompt(overview_text, survey_text, template_filename, survey=None):
    today = datetime.now()
    today_str = f"{today.year}年{today.month}月{today.day}日"
    year_str = str(today.year)

    # 计算实施日期：与手册保持一致（按认证准备周期倒推），确保体系文件时间基准统一
    if survey:
        impl_date_str, impl_date_desc = calculate_implementation_date(survey)
    else:
        impl_date_str = today_str
        impl_date_desc = "未提供调研原始数据，使用当前日期"

    system = (
        "你是程序文件智能生成助手。你会收到一份程序文件模板的结构概览（带段落索引P#、表格T#.R#、页眉H#、页脚F#）"
        "和用户填写的体系调研数据。你的任务是根据调研数据，决定模板中哪些位置需要修改，逐条输出修改方案。\n\n"
        "【输出格式 - 极其重要】\n"
        "每条修改方案单独输出为一行 JSON 对象（NDJSON 格式），不要包裹在数组里，不要输出任何其他文字。"
        "每生成一条就立即输出一行，不要等所有方案想完再一起输出。输出完所有方案后，最后一行写：===END===\n\n"
        "示例输出（每行一个 JSON，逐行输出）：\n"
        '{"type":"global_replace","old":"AAA企业","new":"诸暨正和金属有限公司","reason":"整体替换公司名"}\n'
        '{"type":"paragraph","index":5,"new_text":"本程序适用于XXX公司所有...","reason":"替换适用范围"}\n'
        '{"type":"table_cell","table":0,"row":2,"col":3,"new_text":"王大明","reason":"填入总经理姓名"}\n'
        '===END===\n\n'
        "修改类型说明：\n"
        "- paragraph: 把段落 P#index 整段替换为 new_text（保留段落格式）\n"
        "- table_cell: 把表格 T#table 第 row 行第 col 列单元格替换为 new_text\n"
        "- global_replace: 在全文（段落+表格+页眉页脚）中把 old 替换为 new\n"
        "- header_replace: 仅在页眉页脚中把 old 替换为 new\n\n"
        "关键规则：\n"
        "1. 公司名必须整体替换：模板里的 AAA企业、AAA 等都要整体替换为用户填写的公司名。\n"
        "2. 文件编号：模板里的文件编号（如 AAA-XX-QP-XX）中的 AAA 替换为合适的公司简称，保留编号结构不变。\n"
        "3. 部门名称：根据模板所属部门，填入调研数据中对应的部门负责人姓名。\n"
        "4. 公司宗旨/质量方针/质量目标：如模板中有相关章节，用调研数据替换。\n"
        "5. 实施日期（重要——已按认证准备周期倒推计算，与质量手册保持一致）：程序文件应与手册同步生效，\n"
        "   根据证书类型和第一阶段开始时间倒推得出实施日期为 " + impl_date_str + "。\n"
        "   【计算依据】" + impl_date_desc + "。\n"
        "   模板里所有\"X年X月X日起实施\"、\"发布日期\"、\"生效日期\"、\"实施日期\"、\"制订日期\"、\"审查日期\"、\"批准日期\"等表达程序文件生效时间的日期，\n"
        "   都要替换为 " + impl_date_str + "。\n"
        "   注意：\"计划取得认证证书日期\"不要改，那是认证目标日期，不是程序文件的日期。\n"
        "   程序文件的日期是体系开始运行的日期，必须与手册一致，早于认证日期。\n"
        "6. 人名：总经理/管理者代表/贯标办主任等姓名填入对应签字位置。\n"
        "7. 不要修改 IATF16949/ISO9001 标准条款内容。\n"
        "8. 如果调研数据某字段为空，跳过对应修改。\n"
        "9. 修改方案要全面，覆盖所有该改的位置。\n"
        "10. new_text 中的换行用 \\n 转义。\n"
        "11. 想到一条就立即输出一条。\n"
        "12. 保留模板的结构完整性（但必须替换公司名和敏感信息）：\n"
        "    - 不要修改标题的章节编号（如 1.1、5.2.1 等数字编号）\n"
        "    - 不要修改程序文件的步骤编号（如 5.1、5.2、5.3 等）\n"
        "    - 不要删除或重排模板的章节结构\n"
        "    - 不要修改表格的行列结构（只替换单元格内容）\n"
        "    - 保留模板中的所有标准流程步骤的描述原文不动\n"
        "    【但是】以下内容必须替换，不属于'结构完整性'保护范围：\n"
        "    - 模板中所有出现的 AAA、AAA企业 必须替换为用户公司名\n"
        "    - 标题页/封面/页眉中的公司名必须替换\n"
        "    - 文件编号中的 AAA 必须替换（如 AAA-GM-QP-01 → 正和-GM-QP-01）\n"
        "    - 日期必须替换为实施日期（" + impl_date_str + "），不是当前日期\n"
        "    - 人名必须填入对应位置\n"
        "13. 优先使用 global_replace 做简单替换（如公司名、日期、编号中的AAA），只在需要整段重写时用 paragraph。\n"
        "14. 特别是 global_replace 要覆盖所有 AAA 变体：AAA、AAA企业、山东AAA 等，全部替换为用户公司名。\n"
        "15. 【封面表格填写】可以往封面表格（Table 0）的'实施日期/制订/审查/批准'下方的空单元格填写内容：\n"
        "    - 实施日期下方填入实施日期（" + impl_date_str + "）\n"
        "    - 制订/审查/批准下方填入对应人名（从调研数据获取）\n"
        "    - 填写内容简短即可，不要换行\n"
        "16. 【页眉页脚保护】不要删除或清空页眉页脚中的任何内容，只做替换：\n"
        "    - 页眉中的公司名、文件编号、文件名称、版次、类别、页次等必须保留\n"
        "    - 只用 header_replace 替换其中的 AAA/公司名/编号，不要用 paragraph 清空页眉\n"
        "17. 【重要】所有页面都要检查，不能因为内容少就跳过：\n"
        "    - 修改记录页、附件页、任命书等：公司名、人名、日期、编号都要检查\n"
        "    - 空白表格：表头里的公司名、文件编号也要替换\n"
        "18. 【重要】表格内的内容必须检查：\n"
        "    - 表格单元格里的公司名、文件编号、人名、日期都要用 table_cell 替换\n"
        "    - 职能分配表的部门名称要与用户调研数据中的组织架构一致\n"
        "19. 【重要】即使页面内容很少（只有表头或几个字），也要检查是否有公司名/编号/人名/日期需要替换。\n"
    )

    user = (
        f"当前日期：{today_str}\n"
        f"实施日期（倒推）：{impl_date_str}（{impl_date_desc}）\n"
        f"模板文件：{template_filename}\n\n"
        f"{survey_text}\n\n"
        f"{overview_text}\n\n"
        "请分析以上模板和调研数据，逐条输出 NDJSON 格式的修改方案，每条一行，最后输出 ===END===。"
    )

    return system, user


# ===================================================================
# 4. 解析 LLM 修改方案
# ===================================================================

def parse_ndjson_line(line):
    line = line.strip()
    if not line or line == '===END===':
        return None
    if line.startswith('```'):
        line = line.lstrip('`').replace('json', '', 1).strip()
    if line.endswith('```'):
        line = line[:-3].strip()
    if not line.startswith('{') or not line.endswith('}'):
        return None
    try:
        return json.loads(line)
    except json.JSONDecodeError:
        fixed = re.sub(r',\s*([}\]])', r'\1', line)
        try:
            return json.loads(fixed)
        except json.JSONDecodeError:
            return None


# ===================================================================
# 5. 应用修改方案到 docx
# ===================================================================

def set_paragraph_text(p, new_text):
    if new_text is None:
        new_text = ''
    if not p.runs:
        p.add_run(str(new_text))
        return
    p.runs[0].text = str(new_text)
    for r in p.runs[1:]:
        r.text = ''


def replace_text_in_paragraph(p, old, new):
    """替换段落文字，同时保留不受影响 run 的原有格式。

    Word 经常把一个公司名拆到多个 run 中。直接重设整段文字虽然能替换成功，
    但会把加粗、斜体、颜色、字段等格式全部塞进第一个 run。这里按字符区间从
    后向前修改，只把命中的 run 内容改为新文本；未命中的 run 完全不动。
    """
    if not p.runs or not old:
        return False

    old = str(old)
    new = '' if new is None else str(new)
    run_texts = [(run.text or '') for run in p.runs]
    full = ''.join(run_texts)
    if old not in full:
        return False

    # 每个 run 在拼接字符串中的半开区间 [start, end)。空 run 没有字符，
    # 因而不会作为一次匹配的起点或终点，但它本身的格式也不会被改动。
    spans = []
    offset = 0
    for run_index, text in enumerate(run_texts):
        end = offset + len(text)
        spans.append((offset, end, run_index))
        offset = end

    matches = []
    search_from = 0
    while True:
        start = full.find(old, search_from)
        if start < 0:
            break
        matches.append((start, start + len(old)))
        search_from = start + len(old)

    def _run_for_start(char_offset):
        for run_start, run_end, run_index in spans:
            if run_start <= char_offset < run_end:
                return run_start, run_end, run_index
        return None

    def _run_for_end(char_offset):
        for run_start, run_end, run_index in spans:
            if run_start < char_offset <= run_end:
                return run_start, run_end, run_index
        return None

    # 从后向前处理，前面的原始字符偏移不会被后面的替换长度改变。
    for match_start, match_end in reversed(matches):
        start_run = _run_for_start(match_start)
        end_run = _run_for_end(match_end)
        if start_run is None or end_run is None:
            # 正常 run 文本不可能走到这里；保守跳过，避免破坏文档。
            continue

        start_offset, _, start_index = start_run
        end_offset, _, end_index = end_run
        start_text = p.runs[start_index].text or ''
        prefix = start_text[:match_start - start_offset]

        if start_index == end_index:
            suffix = start_text[match_end - start_offset:]
            p.runs[start_index].text = prefix + new + suffix
            continue

        end_text = p.runs[end_index].text or ''
        suffix = end_text[match_end - end_offset:]
        # 新文字继承命中起点 run 的格式；匹配两端之外的文字仍留在各自原 run。
        p.runs[start_index].text = prefix + new
        for run_index in range(start_index + 1, end_index):
            p.runs[run_index].text = ''
        p.runs[end_index].text = suffix

    return True


def replace_text_in_cell(cell, old, new):
    changed = False
    for p in cell.paragraphs:
        if replace_text_in_paragraph(p, old, new):
            changed = True
    for t in cell.tables:
        for row in t.rows:
            for c in row.cells:
                if replace_text_in_cell(c, old, new):
                    changed = True
    return changed


def apply_global_replace(doc, old, new):
    if new is None:
        new = ''
    if not old or old == new:
        return 0
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            if replace_text_in_paragraph(p, old, new):
                count += 1
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                if old in c.text:
                    if replace_text_in_cell(c, old, new):
                        count += 1
    for sec in doc.sections:
        for hf in [sec.header, sec.first_page_header, sec.even_page_header,
                   sec.footer, sec.first_page_footer, sec.even_page_footer]:
            if not hf:
                continue
            for p in hf.paragraphs:
                if old in p.text:
                    if replace_text_in_paragraph(p, old, new):
                        count += 1
            for t in hf.tables:
                for row in t.rows:
                    for c in row.cells:
                        if old in c.text:
                            if replace_text_in_cell(c, old, new):
                                count += 1
    return count


def apply_header_replace(doc, old, new):
    if new is None:
        new = ''
    if not old or old == new:
        return 0
    count = 0
    for sec in doc.sections:
        for hf in [sec.header, sec.first_page_header, sec.even_page_header,
                   sec.footer, sec.first_page_footer, sec.even_page_footer]:
            if not hf:
                continue
            for p in hf.paragraphs:
                if old in p.text:
                    if replace_text_in_paragraph(p, old, new):
                        count += 1
            for t in hf.tables:
                for row in t.rows:
                    for c in row.cells:
                        if old in c.text:
                            if replace_text_in_cell(c, old, new):
                                count += 1
    return count


def apply_paragraph_replace(doc, index, new_text):
    if index < 0 or index >= len(doc.paragraphs):
        print(f"[WARN] 段落索引 {index} 越界")
        return False
    p = doc.paragraphs[index]
    # 【节分界段落保护】如果段落含 sectPr（节分界符），禁止修改内容
    # 节分界段落必须保持空白，填入内容会导致节属性错乱、页眉引用丢失、空白页
    from docx.oxml.ns import qn
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
        print(f"[INFO] 跳过节分界段落 P{index}（含 sectPr，禁止修改）")
        return False
    set_paragraph_text(p, new_text)
    return True


def apply_table_cell_replace(doc, table_idx, row_idx, col_idx, new_text):
    if table_idx < 0 or table_idx >= len(doc.tables):
        print(f"[WARN] 表格索引 {table_idx} 越界")
        return False
    t = doc.tables[table_idx]
    if row_idx < 0 or row_idx >= len(t.rows):
        print(f"[WARN] 行索引 {row_idx} 越界")
        return False
    row = t.rows[row_idx]
    if col_idx < 0 or col_idx >= len(row.cells):
        print(f"[WARN] 列索引 {col_idx} 越界")
        return False
    cell = row.cells[col_idx]
    if cell.paragraphs:
        set_paragraph_text(cell.paragraphs[0], new_text)
        for p in cell.paragraphs[1:]:
            for r in list(p.runs):
                r.text = ''
    else:
        cell.text = str(new_text)
    return True


def apply_modifications(doc, modifications):
    stats = {
        'paragraph': 0, 'table_cell': 0, 'global_replace': 0,
        'header_replace': 0, 'unknown': 0, 'failed': 0,
    }

    # AI 有时会先给出 AAA -> 新公司名，再给 AAA企业 -> 新公司名。若按原顺序
    # 执行，前一次会把后一次的目标改成“新公司名企业”。所有全文/页眉页脚替换
    # 均按旧文本长度从长到短执行，避免短文本抢先吞掉长文本。其余修改仍保持原序。
    replacement_mods = []
    other_mods = []
    for i, mod in enumerate(modifications):
        if isinstance(mod, dict) and mod.get('type') in ('global_replace', 'header_replace'):
            replacement_mods.append((i, mod))
        else:
            other_mods.append((i, mod))

    def _old_text_length(item):
        old = item[1].get('old', '')
        return len(old) if isinstance(old, str) else len(str(old or ''))

    ordered_mods = sorted(
        replacement_mods,
        key=lambda item: (-_old_text_length(item), item[0])
    ) + other_mods

    for i, mod in ordered_mods:
        try:
            mod_type = mod.get('type', '')
            reason = (mod.get('reason', '') or '')[:80]
            if mod_type == 'paragraph':
                idx = int(mod.get('index', -1))
                new_text = mod.get('new_text', '')
                if apply_paragraph_replace(doc, idx, new_text):
                    stats['paragraph'] += 1
                    print(f"  [P{idx}] OK {reason}")
                else:
                    stats['failed'] += 1
            elif mod_type == 'table_cell':
                ti = int(mod.get('table', -1))
                ri = int(mod.get('row', -1))
                ci = int(mod.get('col', -1))
                new_text = mod.get('new_text', '')
                if apply_table_cell_replace(doc, ti, ri, ci, new_text):
                    stats['table_cell'] += 1
                    print(f"  [T{ti}.R{ri}.C{ci}] OK {reason}")
                else:
                    stats['failed'] += 1
            elif mod_type == 'global_replace':
                old = mod.get('old', '')
                new = mod.get('new', '')
                n = apply_global_replace(doc, old, new)
                stats['global_replace'] += n
                print(f"  [G] '{old}' -> '{new}' ({n} 处) OK {reason}")
            elif mod_type == 'header_replace':
                old = mod.get('old', '')
                new = mod.get('new', '')
                n = apply_header_replace(doc, old, new)
                stats['header_replace'] += n
                print(f"  [H] '{old}' -> '{new}' ({n} 处) OK {reason}")
            else:
                stats['unknown'] += 1
        except Exception as e:
            stats['failed'] += 1
            print(f"  [ERROR] 修改 #{i} 失败: {e}")
    return stats


def remove_even_page_headers_footers(doc):
    """兼容保留的无操作函数：不再删除奇偶页页眉/页脚引用。

    质量体系文件常用不同的偶数页页眉或页脚。过去为了规避个别模板的空白页而
    删除 ``even`` 引用，会静默破坏所有正常的镜像页版式。调用方可以继续调用
    本函数，但它现在保证不修改文档并始终返回 0。
    """
    del doc
    return 0
